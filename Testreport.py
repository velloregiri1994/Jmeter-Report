#!/usr/bin/env python3
"""
------------------------------------------------------------
Optimized Test Completion Report Generator (Improved Version)
------------------------------------------------------------
Author: Giri Reddy (Enhanced by M365 Copilot)
Description:
    - Parses JMeter CSV files
    - Parses Sysdig CPU/Mem JSON files
    - Generates performance graphs
    - Inserts tables/images into Word template
    - Outputs fully automated Test Completion Report
------------------------------------------------------------
"""

import argparse
import csv
import json
import logging
import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Any

from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.dates as mdates

import requests
import urllib3

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

# ---------------------------------------------------------
# LOGGING CONFIGURATION
# ---------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="✅ %(levelname)s: %(message)s"
)

# -------------------------------------------------------------------------
# SYSDIG / PROMETHEUS CONFIG (hardcoded - change here as needed)
# -------------------------------------------------------------------------
SYSDIG_API_URL = "https://app.sysdigcloud.com/prometheus/api/v1"
SYSDIG_API_TOKEN = "bf82af7a-0c05-43ac-94c1-1f3419119d95"  # ← UPDATE YOUR TOKEN HERE
SYSDIG_CLUSTER = "stage-eu-central-1-rg8mt"
SYSDIG_NAMESPACE = "boarding-bcs-ppe"
SYSDIG_STEP = "3"

SYSDIG_WORKLOADS = [
    "bcs-kafka-message-handler",
    "bcs-order-management-aggregator-api",
    "bcs-order-manager-publisher-api",
    "bcs-order-manager-service",
    "bcs-outlet-creation-microservice",
    "bcs-product-catalogue-service",
    "order-management-bg-service",
]


# ---------------------------------------------------------
# UTILITY FUNCTIONS
# ---------------------------------------------------------
def format_epoch(ms: Optional[int]) -> str:
    """Convert epoch milliseconds → UTC timestamp."""
    if ms is None:
        return "N/A"
    try:
        # auto detect if input is seconds or ms
        if ms > 1e12:  # milliseconds
            return datetime.utcfromtimestamp(ms / 1000).strftime("%Y-%m-%d %H:%M:%S")
        else:          # seconds
            return datetime.utcfromtimestamp(ms).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return "N/A"


def percentile(values: List[float], p: float) -> float:
    """Correct linear interpolation percentile calculation."""
    if not values:
        return 0.0
    values = sorted(values)
    k = (len(values) - 1) * (p / 100)
    f = int(k)
    c = min(f + 1, len(values) - 1)
    if f == c:
        return values[f]
    return values[f] + (values[c] - values[f]) * (k - f)


def epoch_second(ts_ms: int) -> Optional[int]:
    """Convert epoch milliseconds → seconds."""
    try:
        if ts_ms > 1e12:
            return ts_ms // 1000       # ms → sec
        return ts_ms                   # already in seconds
    except Exception:
        return None


# ---------------------------------------------------------
# JMETER CSV PARSING
# ---------------------------------------------------------
def parse_jmeter_csv(path: Path) -> List[Dict[str, Any]]:
    """Load JMeter CSV → list of Transaction Controller records only."""
    records = []

    with path.open("r", encoding="utf-8", errors="ignore") as f:
        reader = csv.DictReader(f)

        for idx, row in enumerate(reader, start=1):
            try:
                # ✅ Identify Transaction Controller rows
                response_message = row.get("responseMessage", "")

                # ✅ Keep ONLY Transaction Controllers
                if "Number of samples in transaction" not in response_message:
                    continue  # ❌ Skip samplers

                elapsed = float(row.get("elapsed", 0))
                ts = int(row.get("timeStamp", 0)) / 1000  # convert to seconds

                records.append({
                    "ts": ts,
                    "elapsed": elapsed,
                    "label": row.get("label", "UNKNOWN"),
                    "success": str(row.get("success", "true")).lower() == "true",
                })

            except Exception as e:
                logging.warning(f"Skipping malformed row {idx}: {e}")
                continue

    return records


def compute_metrics(records):
    """Compute aggregated per‑label metrics + response time timeseries (elapsed time)."""
    if not records:
        return {"per_label": [], "timeseries": {}}

    # -------------------------------
    # Group by label
    # -------------------------------
    groups = {}
    for r in records:
        label = r["label"]
        groups.setdefault(label, {"values": [], "errors": 0})
        groups[label]["values"].append(r["elapsed"])
        if not r["success"]:
            groups[label]["errors"] += 1

    # -------------------------------
    # Per‑label summary
    # -------------------------------
    per_label = []
    for label, data in groups.items():
        vals = sorted(data["values"])
        if not vals:
            continue
        per_label.append({
            "label": label,
            "count": len(vals),
            "avg": sum(vals) / len(vals),
            "p90": percentile(vals, 90),
            "p95": percentile(vals, 95),
            "error_rate": (data["errors"] / len(vals)) * 100,
        })

    # -------------------------------
    # Determine test start time
    # -------------------------------
    all_secs = [
        epoch_second(r["ts"])
        for r in records
        if r.get("ts") is not None
    ]
    test_start_sec = min(s for s in all_secs if s is not None) if all_secs else 0

    # -------------------------------
    # Timeseries (avg per second, elapsed time)
    # -------------------------------
    label_ts = {}
    for r in records:
        sec = epoch_second(r["ts"])
        if sec is None:
            continue
        label = r["label"]
        label_ts.setdefault(label, {})
        label_ts[label].setdefault(sec, []).append(r["elapsed"])

    final_ts = {
        label: [
            {
                "elapsed_sec": sec - test_start_sec,
                "avg_ms": sum(vals) / len(vals),
            }
            for sec, vals in sorted(sec_map.items())
        ]
        for label, sec_map in label_ts.items()
    }

    return {"per_label": per_label, "timeseries": final_ts}


def get_start_end_times(records: List[Dict[str, Any]]):
    if not records:
        return None, None
    ts = [r["ts"] for r in records]
    return min(ts), max(ts)


def get_total_average(records: List[Dict[str, Any]]) -> float:
    if not records:
        return 0.0
    return sum(r["elapsed"] for r in records) / len(records)


# ---------------------------------------------------------
# SYSDIG JSON PARSING
# ---------------------------------------------------------
def extract_sysdig_series(obj: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """
    Common Sysdig extractor for CPU & Memory.
    ASSUMPTION:
      - Values are already percentages (0–100)
    """
    result = obj.get("data", {}).get("result", [])
    if not result:
        return None

    timestamps = set()
    container_map: Dict[str, Dict[float, float]] = {}

    for entry in result:
        cname = (
            entry.get("metric", {}).get("container")
            or entry.get("metric", {}).get("kube_workload_name")
            or entry.get("metric", {}).get("kube_pod_name")
            or "container"
        )

        container_map.setdefault(cname, {})

        for ts_raw, val_raw in entry.get("values", []):
            try:
                ts = float(ts_raw)
                if ts > 1e12:
                    ts /= 1000  # ms → sec

                val = float(val_raw)  # ✅ already %
            except Exception:
                continue

            container_map[cname][ts] = val
            timestamps.add(ts)

    timestamps = sorted(timestamps)

    return {
        "labels": [str(t) for t in timestamps],
        "series": [
            {
                "name": cname,
                "values": [ts_map.get(t) for t in timestamps],
            }
            for cname, ts_map in container_map.items()
        ],
    }


def build_infra_metrics(paths: List[str]) -> Dict[str, Optional[Dict[str, Any]]]:
    metrics: Dict[str, Optional[Dict[str, Any]]] = {"CPU": None, "Memory": None}
    for p in paths:
        path = Path(p)
        if not path.exists():
            logging.warning(f"Sysdig file missing: {p}")
            continue

        with path.open("r") as f:
            obj = json.load(f)

        pl = p.lower()
        if "cpu" in pl:
            cpu_data = extract_sysdig_series(obj)
            if cpu_data is not None:
                metrics["CPU"] = cpu_data
        if "mem" in pl:
            mem_data = extract_sysdig_series(obj)
            if mem_data is not None:
                metrics["Memory"] = mem_data

    return metrics


# ---------------------------------------------------------
# SYSDIG API / PromQL helpers
# ---------------------------------------------------------

def workloads_regex_from_list(workloads: List[str]) -> str:
    if not workloads:
        return r"(?!)"
    escaped = [re.escape(w).replace(r'\-', r'-') for w in workloads]
    return "^(" + "|".join(escaped) + ")$"


def query_sysdig_range(
    api_url: str,
    api_token: str,
    query: str,
    start: int,
    end: int,
    step: str = SYSDIG_STEP,
) -> List[Dict[str, Any]]:
    headers = {
        "Authorization": f"Bearer {api_token}",
        "Accept": "application/json",
    }
    params = {"query": query, "start": start, "end": end, "step": step}
    if api_url.endswith("/query_range"):
        url = api_url
    else:
        url = f"{api_url}/query_range"

    logging.info(f"Sysdig query_range request: {query} ({start}-{end}) to {url}")

    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    resp = requests.get(url, headers=headers, params=params, timeout=120, verify=False)
    resp.raise_for_status()

    payload = resp.json()
    if payload.get("status") != "success":
        raise RuntimeError("Sysdig API query_range returned non-success")

    return payload["data"].get("result", [])


def prometheus_matrix_to_series(matrix_result: List[Dict[str, Any]], metric_type: str) -> Dict[str, Any]:
    timestamps = set()
    for obj in matrix_result:
        for ts_val in obj.get("values", []):
            ts = float(ts_val[0])
            if ts > 1e12:
                ts /= 1000.0
            timestamps.add(ts)

    sorted_ts = sorted(timestamps)

    series = []
    for obj in matrix_result:
        metric = obj.get("metric", {})
        workload = metric.get("kube_workload_name", "")
        container = metric.get("container", "")
        # Use workload name as it's unique and cleaner
        if workload:
            name = workload
        elif container:
            name = container
        else:
            name = "unknown"
        values_map = {}
        for ts_val in obj.get("values", []):
            ts = float(ts_val[0])
            if ts > 1e12:
                ts /= 1000.0
            try:
                val = float(ts_val[1])
                values_map[ts] = val
            except Exception:
                values_map[ts] = None

        series.append({"name": name, "values": [values_map.get(ts) for ts in sorted_ts]})

    labels = [str(int(ts)) for ts in sorted_ts]
    return {"labels": labels, "series": series}


def discover_workloads_from_sysdig(
    api_url: str,
    api_token: str,
    cluster: str,
    namespace: str,
    expected_workloads: List[str],
) -> List[str]:
    """
    Query Sysdig label API to fetch all workloads in cluster/namespace.
    Filter to keep only those in expected_workloads list.
    Handles optional rollout suffix (e.g., workload-123).
    """
    if not expected_workloads:
        logging.warning("No expected workloads to discover")
        return []

    headers = {
        "Authorization": f"Bearer {api_token}",
        "Accept": "application/json",
    }

    # Query: get all kube_workload_name values in the cluster/namespace
    # Using label_values endpoint (Prometheus-compatible)
    query = f'{{kube_cluster_name="{cluster}",kube_namespace_name="{namespace}"}}'
    params = {"match[]": query}
    
    if api_url.endswith("/query_range"):
        url = api_url.replace("/query_range", "/label/kube_workload_name/values")
    else:
        url = f"{api_url}/label/kube_workload_name/values"

    logging.info(f"Discovering workloads in {cluster}/{namespace}...")

    try:
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        resp = requests.get(
            url,
            headers=headers,
            params=params,
            timeout=120,
            verify=False,
        )
        resp.raise_for_status()

        payload = resp.json()
        if payload.get("status") != "success":
            logging.warning("Sysdig label discovery did not return success; using fallback list")
            return expected_workloads

        all_workloads = payload.get("data", [])
        logging.info(f"Found {len(all_workloads)} workloads in Sysdig")

        # Filter to keep only expected workloads (with or without rollout suffix)
        base_names_set = set(expected_workloads)
        discovered = []

        for wl in all_workloads:
            # Check if this workload matches any base name (ignoring rollout suffix)
            base = re.sub(r'-[0-9]+$', '', wl)  # strip -digits suffix
            if base in base_names_set and wl not in discovered:
                discovered.append(wl)

        if discovered:
            logging.info(f"✅ Discovered {len(discovered)} matching workloads: {discovered}")
            return sorted(discovered)
        else:
            logging.warning("No matching workloads found in Sysdig; using fallback list")
            return expected_workloads

    except Exception as e:
        logging.warning(f"Sysdig discovery failed ({e}); using fallback list")
        return expected_workloads


# ---------------------------------------------------------
# GRAPH GENERATION
# ---------------------------------------------------------

def format_elapsed_time(seconds):
    seconds = max(0, int(seconds))
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"

from matplotlib.ticker import FuncFormatter, MaxNLocator

def generate_rt_graph(metrics, outfile, start_time=None, end_time=None):
    ts_dict = metrics.get("timeseries", {})
    if not ts_dict:
        return

    # ✅ Create figure & axis explicitly (IMPORTANT)
    fig, ax = plt.subplots(figsize=(20, 10))

    for label, ts_list in ts_dict.items():
        times = [t["elapsed_sec"] for t in ts_list]
        values = [t["avg_ms"] for t in ts_list]
        ax.plot(
            times,
            values,
            linewidth=2,
            alpha=0.9,
            label=label,
        )

    # ✅ Titles and labels
    ax.set_title(
        "Average Response Time Over Test Duration (Transaction Controllers Only)",
        fontsize=16,
        fontweight="bold",
    )
    ax.set_xlabel("Elapsed Time (HH:MM:SS)", fontsize=12)
    ax.set_ylabel("Average Response Time (ms)", fontsize=12)

    # ✅ Grid
    ax.grid(True, linestyle="--", alpha=0.4)

    # ✅ X‑axis formatting - improved for readability
    ax.xaxis.set_major_formatter(
        FuncFormatter(lambda x, _: format_elapsed_time(x))
    )
    ax.xaxis.set_major_locator(MaxNLocator(10))  # Reduced ticks for clarity
    ax.tick_params(axis="x", rotation=45, labelsize=10)  # Rotate for better fit

    # ✅ add run timestamp labels when provided
    if start_time and end_time:
        try:
            ax.text(
                0.02,
                0.96,
                f"Start: {start_time}   End: {end_time}",
                transform=ax.transAxes,
                fontsize=10,
                color="dimgray",
                bbox=dict(facecolor="white", alpha=0.6, edgecolor="none"),
            )
        except Exception:
            pass

    # ✅ LEGEND UNDER THE CHART — SAFE VERSION
    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, -0.20),
        ncol=3,
        fontsize=11,
        frameon=False,
        title="Transactions",
        title_fontsize=12,
        handlelength=3,
    )

    # ✅ CRITICAL FIX: adjust bottom BEFORE saving
    fig.subplots_adjust(bottom=0.30)

    # ✅ Save & close explicitly
    fig.savefig(outfile, dpi=180, bbox_inches="tight")
    plt.close(fig)



def extract_series_start_end(series: Dict[str, Any]) -> Dict[str, Dict[str, float]]:
    """Extract start and end values for each series."""
    result = {}
    for s in series.get("series", []):
        values = [v for v in s["values"] if v is not None]
        if values:
            start_val = values[0]
            end_val = values[-1]
            result[s["name"]] = {"start": start_val, "end": end_val}
    return result


def generate_sysdig_chart(
    series: Optional[Dict[str, Any]],
    outfile: Path,
    metric_type: str,  # "CPU" or "Memory"
):
    if not series:
        logging.warning(f"No Sysdig data for {metric_type}")
        return

    if not series.get("series") or not series["series"]:
        logging.warning(f"No series data for {metric_type}")
        return

    times = [datetime.utcfromtimestamp(float(t)) for t in series["labels"]]

    fig, ax = plt.subplots(figsize=(18, 8))

    # Extract start/end values for each series
    start_end_data = extract_series_start_end(series)

    for s in series["series"]:
        values = [v for v in s["values"] if v is not None]
        if values:
            min_val = min(values)
            max_val = max(values)
            logging.info(f"📊 {metric_type} series '{s['name']}': min={min_val:.2f}, max={max_val:.2f}")
        # Create legend label with start/end values
        se_data = start_end_data.get(s["name"], {})
        start = se_data.get("start", 0)
        end = se_data.get("end", 0)
        legend_label = f"{s['name']}: {start:.1f}%-{end:.1f}%"
        ax.plot(times, s["values"], linewidth=2, alpha=0.85, label=legend_label)  # type: ignore

    # ✅ Titles & labels based on metric
    if metric_type == "Memory":
        ax.set_title(f"{metric_type} Utilization Over Time", fontsize=15, fontweight="bold")
        ax.set_ylabel(f"{metric_type} Utilization (%)", fontsize=12)
        ax.set_ylim(0, 100)
    else:
        ax.set_title(f"{metric_type} Utilization Over Time", fontsize=15, fontweight="bold")
        ax.set_ylabel(f"{metric_type} Utilization (%)", fontsize=12)
        ax.set_ylim(0, 100)

    ax.grid(True, linestyle="--", alpha=0.4)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d %H:%M"))  # Include date for clarity
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right", fontsize=9)  # Smaller font, better rotation

    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, -0.25),
        ncol=2,
        fontsize=9,
        frameon=False,
        title="Containers (Start% → End%)",
        title_fontsize=10,
    )

    fig.subplots_adjust(bottom=0.30)
    fig.savefig(outfile, dpi=160, bbox_inches="tight")
    plt.close(fig)


# ---------------------------------------------------------
# DOCX TEMPLATING HELPERS
# ---------------------------------------------------------
def deep_replace_text(element, placeholder, new_text):
    """Search & replace inside Word paragraphs/tables safely."""
    if hasattr(element, "paragraphs"):
        for p in element.paragraphs:
            full = "".join(run.text for run in p.runs)
            if placeholder in full:
                new_full = full.replace(placeholder, new_text)
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_full
                else:
                    p.add_run(new_full)

    if hasattr(element, "tables"):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    deep_replace_text(cell, placeholder, new_text)


def replace_text_safely(doc: Document, placeholder: str, new_text: str):  # type: ignore
    """Utility wrapper for deep replacement."""
    for element in [doc]:
        deep_replace_text(element, placeholder, new_text)


def insert_paragraph_after(paragraph):
    """Insert new paragraph after given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_table_after_paragraph(paragraph, rows, cols):
    """Insert a table after a given paragraph."""
    tbl = OxmlElement("w:tbl")

    tbl_pr = OxmlElement("w:tblPr")
    tbl.append(tbl_pr)

    tbl_grid = OxmlElement("w:tblGrid")
    for _ in range(cols):
        tbl_grid.append(OxmlElement("w:gridCol"))
    tbl.append(tbl_grid)

    for _ in range(rows):
        tr = OxmlElement("w:tr")
        for _ in range(cols):
            tc = OxmlElement("w:tc")
            p = OxmlElement("w:p")
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def find_paragraph_index(doc, heading, allowed_levels=("Heading 1", "Heading 2")):
    target = heading.lower().replace(" ", "")

    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ""

        if (
            style_name in allowed_levels and
            target in p.text.lower().replace(" ", "")
        ):
            return i

    return None
    
def set_table_borders(table: Table):
    """Apply full borders to a Word table."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")

    for border_name in (
        "top", "left", "bottom", "right",
        "insideH", "insideV"
    ):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")      # thickness
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)

    tblPr.append(tblBorders)


def insert_rt_summary_table(doc, heading, load_rows, soak_rows):
    """Insert response time summary table under a heading."""

    idx = find_paragraph_index(doc, heading)
    if idx is None:
        logging.warning(f"Heading not found: {heading}")
        return

    # Anchor paragraph (heading)
    anchor = doc.paragraphs[idx]
    new_p = insert_paragraph_after(anchor)

    # ✅ Create table with TWO header rows and SIX columns
    table = insert_table_after_paragraph(new_p, 2, 6)
    set_table_borders(table)

    # =================================================
    # ✅ STEP 2: BUILD GROUPED HEADER (Load / Soak)
    # =================================================
    hdr1 = table.rows[0].cells
    hdr2 = table.rows[1].cells

    # Row 1 (group headers)
    hdr1[0].text = "#"
    hdr1[1].text = "Transaction Name"
    hdr1[2].text = "Load Test"
    hdr1[4].text = "Soak Test"

    # Merge grouped columns
    hdr1[2].merge(hdr1[3])   # Load Test spans 2 columns
    hdr1[4].merge(hdr1[5])   # Soak Test spans 2 columns

    # Row 2 (sub headers)
    hdr2[2].text = "Avg Response (ms)"
    hdr2[3].text = "90th percentile (ms)"
    hdr2[4].text = "Avg Response (ms)"
    hdr2[5].text = "90th percentile (ms)"

    # =================================================
    # ✅ STEP 3: ADD DATA ROWS
    # =================================================
    soak_lookup = {r["label"]: r for r in soak_rows}

    for i, load in enumerate(load_rows, start=1):
        soak = soak_lookup.get(load["label"], {})
        row = table.add_row().cells

        row[0].text = str(i)
        row[1].text = load["label"]
        row[2].text = f"{load['avg']:.2f}"
        row[3].text = f"{load['p90']:.2f}"
        row[4].text = f"{soak.get('avg', 0):.2f}"
        row[5].text = f"{soak.get('p90', 0):.2f}"


def insert_image(doc, placeholder, path):
    """Replace placeholder with image word-safely."""
    if not os.path.exists(path):
        logging.warning(f"Image not found: {path}")
        return

    for p in doc.paragraphs:
        full = "".join(run.text for run in p.runs)
        if placeholder in full:
            new_full = full.replace(placeholder, "")
            for run in p.runs:
                run.text = ""
            p.runs[0].text = new_full
            p.add_run().add_picture(path, width=Inches(6))
            return


# ---------------------------------------------------------
# BUILD COMPLETE REPORT
# ---------------------------------------------------------
def build_report(
    template: Path,
    output: Path,
    load_metrics: Dict[str, Any],
    soak_metrics: Dict[str, Any],
    load_infra: Dict[str, Any],
    soak_infra: Dict[str, Any],
    load_start: str, load_end: str,
    soak_start: str, soak_end: str,
    load_avg: float, soak_avg: float
):
    doc = Document(str(template))

    insert_rt_summary_table(
        doc,
        "8.1 Response time Summary",
        load_metrics["per_label"],
        soak_metrics["per_label"]
    )

    replace_text_safely(doc, "{{LOAD_START_TIME}}", load_start)
    replace_text_safely(doc, "{{LOAD_END_TIME}}", load_end)
    replace_text_safely(doc, "{{SOAK_START_TIME}}", soak_start)
    replace_text_safely(doc, "{{SOAK_END_TIME}}", soak_end)

    replace_text_safely(doc, "{{LOAD_TOTAL_AVG}}", f"{load_avg:.2f}")
    replace_text_safely(doc, "{{SOAK_TOTAL_AVG}}", f"{soak_avg:.2f}")

    insert_image(doc, "{{RT_LOAD_GRAPH}}", "output/rt_load.png")
    insert_image(doc, "{{RT_SOAK_GRAPH}}", "output/rt_soak.png")

    insert_image(doc, "{{CPU_LOAD_GRAPH}}", "output/cpu_load.png")
    insert_image(doc, "{{MEMORY_LOAD_GRAPH}}", "output/memory_load.png")

    insert_image(doc, "{{CPU_SOAK_GRAPH}}", "output/cpu_soak.png")
    insert_image(doc, "{{MEMORY_SOAK_GRAPH}}", "output/memory_soak.png")

    doc.save(str(output))
    logging.info(f"✅ Report generated: {output}")


# ---------------------------------------------------------
# MAIN
# ---------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Generate Test Completion Report")
    parser.add_argument("--load-test", required=True, help="Path to JMeter load test CSV")
    parser.add_argument("--soak-test", required=True, help="Path to JMeter soak test CSV")
    parser.add_argument("--template", required=True, help="Path to Word template")
    parser.add_argument("--output", required=True, help="Output Word document path")

    args = parser.parse_args()

    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)

    load_records = parse_jmeter_csv(Path(args.load_test))
    soak_records = parse_jmeter_csv(Path(args.soak_test))

    load_metrics = compute_metrics(load_records)
    soak_metrics = compute_metrics(soak_records)

    load_avg = get_total_average(load_records)
    soak_avg = get_total_average(soak_records)

    ls, le = get_start_end_times(load_records)
    ss, se = get_start_end_times(soak_records)

    load_start = format_epoch(ls)
    load_end = format_epoch(le)
    soak_start = format_epoch(ss)
    soak_end = format_epoch(se)

    sysdig_api_url = SYSDIG_API_URL.rstrip("/")
    sysdig_api_token = SYSDIG_API_TOKEN

    if not sysdig_api_token:
        raise RuntimeError("Sysdig API token is required. Update SYSDIG_API_TOKEN in the script.")

    # Auto-discover active workloads from Sysdig (filter to expected list)
    logging.info("🔍 Discovering workloads from Sysdig API...")
    discovered_workloads = discover_workloads_from_sysdig(
        sysdig_api_url,
        sysdig_api_token,
        SYSDIG_CLUSTER,
        SYSDIG_NAMESPACE,
        SYSDIG_WORKLOADS,  # use fixed list as filter
    )
    workload_names = discovered_workloads
    workload_regex = workloads_regex_from_list(workload_names)

    cpu_query = (
        f'sysdig_container_cpu_quota_used_percent{{'
        f'kube_cluster_name="{SYSDIG_CLUSTER}",'
        f'kube_namespace_name="{SYSDIG_NAMESPACE}",'
        f'kube_workload_name=~"{workload_regex}"'
        f'}}'
    )
    mem_query = (
        f'sysdig_container_memory_limit_used_percent{{'
        f'kube_cluster_name="{SYSDIG_CLUSTER}",'
        f'kube_namespace_name="{SYSDIG_NAMESPACE}",'
        f'kube_workload_name=~"{workload_regex}"'
        f'}}'
    )

    generate_rt_graph(load_metrics, Path("output/rt_load.png"), load_start, load_end)
    generate_rt_graph(soak_metrics, Path("output/rt_soak.png"), soak_start, soak_end)

    # Query Sysdig for load window, if available
    if ls is not None and le is not None:
        try:
            logging.info("📡 Running Sysdig CPU/Mem queries for load window")
            cpu_matrix_load = query_sysdig_range(sysdig_api_url, sysdig_api_token, cpu_query, ls, le, SYSDIG_STEP)
            mem_matrix_load = query_sysdig_range(sysdig_api_url, sysdig_api_token, mem_query, ls, le, SYSDIG_STEP)
            cpu_series_load = prometheus_matrix_to_series(cpu_matrix_load, "CPU")
            mem_series_load = prometheus_matrix_to_series(mem_matrix_load, "Memory")
            logging.info(f"📊 CPU load series: {len(cpu_series_load['series'])} series with {len(cpu_series_load['labels'])} data points")
            logging.info(f"📊 Memory load series: {len(mem_series_load['series'])} series with {len(mem_series_load['labels'])} data points")
            generate_sysdig_chart(cpu_series_load, Path("output/cpu_load.png"), "CPU Usage - Load Test")
            generate_sysdig_chart(mem_series_load, Path("output/memory_load.png"), "Memory Usage - Load Test")
        except requests.HTTPError as e:
            logging.warning(f"⚠️ Failed to query Sysdig for load window: {e}")
            logging.warning(f"Response status: {e.response.status_code}, text: {e.response.text[:500]}")
    else:
        logging.warning("⚠️ Load window not found; skipping Sysdig load charts")

    # Query Sysdig for soak window, if available
    if ss is not None and se is not None:
        try:
            logging.info("📡 Running Sysdig CPU/Mem queries for soak window")
            cpu_matrix_soak = query_sysdig_range(sysdig_api_url, sysdig_api_token, cpu_query, ss, se, SYSDIG_STEP)
            mem_matrix_soak = query_sysdig_range(sysdig_api_url, sysdig_api_token, mem_query, ss, se, SYSDIG_STEP)
            cpu_series_soak = prometheus_matrix_to_series(cpu_matrix_soak, "CPU")
            mem_series_soak = prometheus_matrix_to_series(mem_matrix_soak, "Memory")
            logging.info(f"📊 CPU soak series: {len(cpu_series_soak['series'])} series with {len(cpu_series_soak['labels'])} data points")
            logging.info(f"📊 Memory soak series: {len(mem_series_soak['series'])} series with {len(mem_series_soak['labels'])} data points")
            generate_sysdig_chart(cpu_series_soak, Path("output/cpu_soak.png"), "CPU Usage - Soak Test")
            generate_sysdig_chart(mem_series_soak, Path("output/memory_soak.png"), "Memory Usage - Soak Test")
        except requests.HTTPError as e:
            logging.warning(f"⚠️ Failed to query Sysdig for soak window: {e}")
            logging.warning(f"Response status: {e.response.status_code}, text: {e.response.text[:500]}")
    else:
        logging.warning("⚠️ Soak window not found; skipping Sysdig soak charts")

    build_report(
        Path(args.template),
        Path(args.output),
        load_metrics,
        soak_metrics,
        {},
        {},
        load_start, load_end,
        soak_start, soak_end,
        load_avg, soak_avg,
    )

    logging.info("✅ All tasks completed successfully")


if __name__ == "__main__":
    main()